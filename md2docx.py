# ==============================================================================
#  Markdown to Docx / PDF Converter v1.0.0
# ==============================================================================
#  [Developed by]
#    JINDE (NIK Co., Ltd.) / 株式会社ニッキ 神出
#
#  [License]
#    MIT License
#
#  [Terms of Use]
#    Free to use for any purpose, provided that the copyright notice is retained.
#    (著作権表示さえあれば、何に使ってもOKです)
#
#  [Usage Examples]
#    # Convert to Docx
#    python3 md2docx.py input.md
#
#    # Convert to PDF (Requires: sudo apt install python3-reportlab)
#    python3 md2docx.py input.md --pdf
#    python3 md2docx.py input.md output.pdf
# ==============================================================================

import re
import os
import base64
import urllib.request
import io
import zlib
import html
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse

VERSION = "1.0.0"
CREDITS = "Developed by JINDE (NIK Co., Ltd.) / 株式会社ニッキ 神出"

# Try to import ReportLab for PDF support
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table as RLTable, TableStyle, PageBreak, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import inch, mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.graphics.shapes import Drawing, Line
    HAS_REPORTLAB = True

except ImportError:
    HAS_REPORTLAB = False

def load_config():
    config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')
    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Configuration file (config.json) not found at: {config_path}")
    
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        raise RuntimeError(f"Error loading config.json: {e}")

def add_table_borders(table):
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '4')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        
        tcBorders.append(top)
        tcBorders.append(left)
        tcBorders.append(bottom)
        tcBorders.append(right)
        tcPr.append(tcBorders)

class DocxWriter:
    def __init__(self):
        self.doc = Document()
        self.config = load_config()
        fonts = self.config['fonts']

        # スタイルの調整 (ゴシック系フォントの設定)
        style = self.doc.styles['Normal']
        font = style.font
        font.name = fonts['normal']['name']
        font.element.rPr.rFonts.set(qn('w:eastAsia'), fonts['normal']['eastAsia'])
        font.size = Pt(fonts['normal']['size'])
        
        # 見出しのサイズ設定
        heading_sizes = fonts['heading']['sizes']
        h_color = fonts['heading']['colors']
        h_name = fonts['heading']['name']
        h_eastAsia = fonts['heading']['eastAsia']

        for i, size in enumerate(heading_sizes, 1):
            if f'Heading {i}' in self.doc.styles:
                h_style = self.doc.styles[f'Heading {i}']
                h_font = h_style.font
                h_font.name = h_name
                h_font.element.rPr.rFonts.set(qn('w:eastAsia'), h_eastAsia)
                h_font.size = Pt(size)
                # 見出しを目立たせる色設定
                h_font.color.rgb = RGBColor(*h_color)

                # スタイルレベルでの自動改ページ設定は無効化 (先頭で改ページされるのを防ぐため)
                # 個別の add_heading メソッド内で制御する
                h_style.paragraph_format.page_break_before = False

    def process_inline_formatting(self, paragraph, text):
        """行内のフォーマット（**太字**）を解析してParagraphに追加"""
        text = text.replace('<br>', '\n')
        parts = re.split(r'(\*\*.*?\*\*)', text)
        bold_color = self.config['fonts']['bold']['colors']
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                # 太字を色変更して通常の太字と区別
                run.font.color.rgb = RGBColor(*bold_color)
            else:
                if part:
                    paragraph.add_run(part)

    def add_heading(self, text, level):
        # 文書の先頭かどうかを判定 (段落も表もまだ無い場合)
        is_first = (len(self.doc.paragraphs) == 0 and len(self.doc.tables) == 0)

        h = self.doc.add_heading('', level=level)
        self.process_inline_formatting(h, text)

        # 見出しの下に線を追加 (H1, H2のみ)
        if level <= 2:
            pPr = h._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6') # 0.75pt
            bottom.set(qn('w:space'), '1')
            
            # 見出し色を取得して線に適用
            h_rgb = self.config['fonts']['heading']['colors']
            hex_color = '{:02x}{:02x}{:02x}'.format(*h_rgb)
            bottom.set(qn('w:color'), hex_color)
            
            pBdr.append(bottom)
            pPr.append(pBdr)

        # 改ページ設定の読み込み (デフォルトはレベル2まで)
        break_level = self.config['fonts']['heading'].get('page_break_level', 2)

        # 設定レベル以下、かつ文書の先頭でない場合は「段落前で改ページ」を有効にする
        if level <= break_level and not is_first:
            h.paragraph_format.page_break_before = True

    def add_paragraph(self, text, style=None):
        p = self.doc.add_paragraph(style=style)
        self.process_inline_formatting(p, text)
        if style == 'No Spacing': # Code block style adjustments
            code_cfg = self.config['fonts']['code']
            for run in p.runs:
                run.font.name = code_cfg['docx_name']
                run.font.color.rgb = RGBColor(*code_cfg['colors'])
        return p

    def add_quote(self, text):
        p = self.doc.add_paragraph()
        self.process_inline_formatting(p, text)
        for run in p.runs:
            run.italic = True

    def add_image(self, image_data, width_inches=5):
        try:
            self.doc.add_picture(image_data, width=Inches(width_inches))
        except Exception:
            self.doc.add_paragraph("[Image insertion failed]")

    def add_table(self, rows, cols, data):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for r_idx, row_data in enumerate(data):
            for c_idx, text in enumerate(row_data):
                if c_idx < cols:
                    cell = table.cell(r_idx, c_idx)
                    p = cell.paragraphs[0]
                    p.text = ""
                    self.process_inline_formatting(p, text)
        add_table_borders(table) # Optional custom borders
        self.doc.add_paragraph()

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        print(f"Word file created: {path}")

    def add_code_block(self, lines):
        if not lines:
            return
            
        # 1行1列の表を作成して枠を作る
        table = self.doc.add_table(rows=1, cols=1)
        
        # 表全体のインデント設定
        tblPr = table._tbl.tblPr
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '400') # 20pt (1pt = 20dxa)
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        add_table_borders(table)
        
        cell = table.cell(0, 0)
        # デフォルトの段落を取得
        p = cell.paragraphs[0]
        # 上下の余白を少し調整
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        # フォント設定
        code_cfg = self.config['fonts']['code']
        
        # 行ごとにRunを追加して強制改行を入れる
        for i, line in enumerate(lines):
            # Tabをスペースに置換
            text = line.replace('\t', '    ')
            
            # DocxではHTMLエスケープや&nbsp;置換は不要
            run = p.add_run(text)
            run.font.name = code_cfg['docx_name']
            run.font.color.rgb = RGBColor(*code_cfg['colors'])
            run.font.size = Pt(code_cfg['size'])
            
            if i < len(lines) - 1:
                run.add_break()

class PdfWriter:
    def __init__(self):
        if not HAS_REPORTLAB:
            print("\n" + "="*60)
            print(" ERROR: ReportLab library is missing.")
            print(" To generate PDF, please install 'reportlab'.")
            print(" For Ubuntu 24.04+ (Managed Environment), run:")
            print("     sudo apt update && sudo apt install -y python3-reportlab")
            print(" OR (using pip):")
            print("     pip install reportlab --break-system-packages")
            print("="*60 + "\n")
            raise ImportError("ReportLab is not installed. Cannot generate PDF.")
        self.config = load_config()
        self.story = []
        self.styles = getSampleStyleSheet()
        self.setup_fonts()
        
        # Define styles
        fonts = self.config['fonts']
        
        self.normal_style = self.styles['Normal']
        self.normal_style.fontName = fonts['normal']['pdf_name']
        self.normal_style.fontSize = fonts['normal']['size']
        self.normal_style.leading = fonts['normal']['size'] * 1.4
        
        c_rgb = fonts['code']['colors']
        c_color = colors.Color(c_rgb[0]/255.0, c_rgb[1]/255.0, c_rgb[2]/255.0)

        self.code_style = ParagraphStyle('Code', parent=self.styles['Normal'], 
                                         fontName=fonts['code']['pdf_name'], 
                                         fontSize=fonts['code']['size'], 
                                         textColor=c_color)
        self.quote_style = ParagraphStyle('Quote', parent=self.styles['Normal'], fontName='JapaneseFont', leftIndent=20, textColor=colors.darkgrey)

    def add_code_block(self, lines):
        if not lines:
            return

        formatted_lines = []
        for line in lines:
            line = line.replace('\t', '    ') 
            line = html.escape(line)
            line = line.replace(' ', '&nbsp;') 
            formatted_lines.append(line)
        
        full_text = '<br/>'.join(formatted_lines)
        
        c_rgb = self.config['fonts']['code']['colors']
        c_color = colors.Color(c_rgb[0]/255.0, c_rgb[1]/255.0, c_rgb[2]/255.0)
        
        block_style = ParagraphStyle(
            'CodeBlock',
            parent=self.styles['Normal'],
            fontName=self.config['fonts']['code']['pdf_name'],
            fontSize=self.config['fonts']['code']['size'],
            textColor=c_color,
            leading=self.config['fonts']['code']['size'] * 1.4,
            leftIndent=20,    # 左インデント追加
            rightIndent=20,   # 右インデント追加
            borderWidth=0.5,
            borderColor=colors.black,
            borderPadding=6,
            backColor=colors.whitesmoke,
            splitLongWords=False
        )
        
        # 枠線の外側に空白を追加（前）
        self.story.append(Spacer(1, 0.15*inch))
        self.story.append(Paragraph(full_text, block_style))
        # 枠線の外側に空白を追加（後）
        self.story.append(Spacer(1, 0.15*inch))

    def setup_fonts(self):
        # 日本語フォントの検索と登録
        font_paths = self.config.get('pdf_font_paths', [])
        
        registered = False
        from reportlab.lib.fonts import addMapping

        for p in font_paths:
            if not os.path.exists(p):
                continue
            
            try:
                # フォントの読み込み試行（TTC対応）
                # 太字(Bold)・斜体(Italic)も同じフォントファイルを割り当てて、エラー("Can't map")を回避する
                if p.endswith('.ttc'):
                    pdfmetrics.registerFont(TTFont('JapaneseFont', p, subfontIndex=0))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-Bold', p, subfontIndex=0))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-Italic', p, subfontIndex=0))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-BoldItalic', p, subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont('JapaneseFont', p))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-Bold', p))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-Italic', p))
                    pdfmetrics.registerFont(TTFont('JapaneseFont-BoldItalic', p))
                
                # スタイルマッピング (Normal, Bold, Italic, BoldItalic)
                # これにより <b>タグなどが来てもクラッシュせず Regular フォントで表示される
                addMapping('JapaneseFont', 0, 0, 'JapaneseFont')            # Normal
                addMapping('JapaneseFont', 0, 1, 'JapaneseFont-Italic')     # Italic
                addMapping('JapaneseFont', 1, 0, 'JapaneseFont-Bold')       # Bold
                addMapping('JapaneseFont', 1, 1, 'JapaneseFont-BoldItalic') # BoldItalic
                
                registered = True
                print(f"Loaded font: {p}")
                break
            except Exception as e:
                print(f"Font loading error ({p}): {e}")
                # 失敗したら次の候補へ
                continue
        
        if not registered:
            self.setup_fallback_font()

    def setup_fallback_font(self):
        print("Warning: Japanese font not found. Characters may not display correctly.")
        from reportlab.lib.fonts import addMapping
        # 日本語が出なくてもクラッシュを防ぐため、システムにある英字フォントで代用登録を試みる
        fallback_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/freefont/FreeSans.ttf"
        ]
        for p in fallback_paths:
            if os.path.exists(p):
                try:
                    pdfmetrics.registerFont(TTFont('JapaneseFont', p))
                    # Fallbackでもマッピングしておく
                    addMapping('JapaneseFont', 0, 0, 'JapaneseFont')
                    addMapping('JapaneseFont', 1, 0, 'JapaneseFont')
                    return
                except:
                    continue

    def _format_text(self, text):
        """Markdownの太字などをReportLabのXMLタグに変換"""
        text = html.escape(text)
        text = text.replace('&lt;br&gt;', '<br/>')
        
        # タブ文字の置換はここで行わず、呼び出し元で制御する

        # **Bold** -> <b>Bold</b> (着色)
        b_rgb = self.config['fonts']['bold']['colors']
        hex_color = '#{:02x}{:02x}{:02x}'.format(*b_rgb)
        text = re.sub(r'\*\*(.*?)\*\*', f'<font color="{hex_color}"><b>\\1</b></font>', text)
        return text

    def add_heading(self, text, level):
        h_cfg = self.config['fonts']['heading']
        sizes = h_cfg['sizes']
        # level is 1-based, sizes is 0-based
        size = sizes[level-1] if 0 < level <= len(sizes) else 12
        
        # 改ページ設定の読み込み (デフォルトはレベル2まで)
        break_level = h_cfg.get('page_break_level', 2)

        # すでに要素が存在する場合のみ改ページする (先頭ページでの改ページ回避)
        use_page_break = (level <= break_level and len(self.story) > 0)
        
        h_rgb = h_cfg['colors']
        h_color = colors.Color(h_rgb[0]/255.0, h_rgb[1]/255.0, h_rgb[2]/255.0)

        # 見出しを目立たせる
        style = ParagraphStyle(
            f'H{level}', 
            parent=self.styles['Normal'], 
            fontName='JapaneseFont', 
            fontSize=size, 
            leading=size*1.2, 
            spaceAfter=6, 
            textColor=h_color,
            pageBreakBefore=use_page_break
        )
        self.story.append(Paragraph(self._format_text(text), style))

        # 見出しの下に線を追加 (H1, H2のみ)
        if level <= 2:
            # 描画エリアの幅 (A4 width - margins)
            width = A4[0] - 144
            d = Drawing(width, 10)
            d.add(Line(0, 5, width, 5, strokeColor=h_color, strokeWidth=1))
            self.story.append(d)
            self.story.append(Spacer(1, 0.1*inch))

    def add_paragraph(self, text, style=None):
        pst = self.normal_style
        is_code = False
        if style == 'No Spacing' or style == 'Code': # Mapping 'No Spacing' used for code mainly
            pst = self.code_style
            is_code = True
        elif style == 'List Bullet':
            text = f"• {text}"
            pst = ParagraphStyle('Bullet', parent=self.normal_style, leftIndent=10)
        elif style == 'List Number':
            pst = ParagraphStyle('Number', parent=self.normal_style, leftIndent=10)
        
        # コードブロックの場合のみタブとスペースを置換
        if is_code:
            text = text.replace('\t', '    ') # タブをスペース4つに
            formatted_text = self._format_text(text)
            formatted_text = formatted_text.replace(' ', '&nbsp;') # スペースを維持用タグに
        else:
            formatted_text = self._format_text(text)

        self.story.append(Paragraph(formatted_text, pst))
        if style != 'No Spacing':
            self.story.append(Spacer(1, 0.1*inch))

    def add_quote(self, text):
        self.story.append(Paragraph(self._format_text(text), self.quote_style))
        self.story.append(Spacer(1, 0.1*inch))

    def add_image(self, image_data, width_inches=5):
        # image_data can be path (str) or bytes stream
        try:
            img = RLImage(image_data)
            # Resize logic simplistic
            aspect = img.imageHeight / float(img.imageWidth)
            disp_width = width_inches * inch
            disp_height = disp_width * aspect
            img.drawHeight = disp_height
            img.drawWidth = disp_width
            self.story.append(img)
            self.story.append(Spacer(1, 0.2*inch))
        except Exception as e:
            print(f"Image error: {e}")

    def add_table(self, rows, cols, data):
        # formatted data
        tbl_data = []
        for row in data:
            row_cells = []
            for cell_text in row:
                row_cells.append(Paragraph(self._format_text(cell_text), self.normal_style))
            tbl_data.append(row_cells)
        
        t = RLTable(tbl_data, colWidths=[(6.0/cols)*inch]*cols)
        t.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        self.story.append(t)
        self.story.append(Spacer(1, 0.2*inch))

    def add_page_break(self):
        self.story.append(PageBreak())

    def save(self, path):
        doc = SimpleDocTemplate(path, pagesize=A4,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
        doc.build(self.story)
        print(f"PDF file created: {path}")

def convert_markdown(md_path, output_path, writer_class=DocxWriter):
    writer = writer_class()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 末尾が改行で終わっていない場合にバッファ（表など）がフラッシュされない問題への対策
    # 強制的に空行を追加することで、ループの最後でelseブロック（フラッシュ処理）に入るようにする
    lines.append("")

    in_code_block = False
    in_mermaid_block = False
    mermaid_lines = []
    table_lines = []
    code_buffer = [] # コードブロック用バッファ
    
    for line in lines:
        # strip() してしまうとコード内のインデントが消えるため、
        # コンテンツ用(raw_line)と判定用(stripped_line)を分ける
        raw_line = line.rstrip('\r\n') 
        stripped_line = raw_line.strip()
        
        # コードブロックの処理
        if stripped_line.startswith('```'):
            if stripped_line.startswith('```mermaid'):
                in_mermaid_block = True
                mermaid_lines = []
                continue
            
            if in_mermaid_block:
                in_mermaid_block = False
                if mermaid_lines:
                    try:
                        mm_code = '\n'.join(mermaid_lines)
                        mm_bytes = mm_code.encode('utf8')
                        mm_compressed = zlib.compress(mm_bytes)
                        mm_b64 = base64.urlsafe_b64encode(mm_compressed).decode('ascii')
                        url = f"https://kroki.io/mermaid/png/{mm_b64}"
                        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req) as response:
                            image_data = io.BytesIO(response.read())
                            writer.add_image(image_data, width_inches=6)
                    except Exception as e:
                        print(f"Mermaid conversion error: {e}")
                        # フォールバックでコードブロックとして出力
                        writer.add_code_block(mermaid_lines)
                continue

            if in_code_block:
                # コードブロック終了：バッファを書き出し
                writer.add_code_block(code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                # コードブロック開始
                in_code_block = True
            continue
            
        if in_mermaid_block:
            mermaid_lines.append(raw_line)
            continue
            
        if in_code_block:
            # バッファに追加
            code_buffer.append(raw_line)
            continue

        # 表の処理
        if stripped_line.startswith('|'):
            table_lines.append(stripped_line)
            continue
        else:
            if table_lines:
                wb_rows = [r for r in table_lines if '---' not in r]
                if wb_rows:
                    cols = len([c for c in wb_rows[0].split('|') if c])
                    data = []
                    for row_text in wb_rows:
                        row_cells = [c.strip() for c in row_text.split('|') if c]
                        # Pad with empty strings if necessary
                        while len(row_cells) < cols: row_cells.append("")
                        data.append(row_cells[:cols])
                    writer.add_table(len(wb_rows), cols, data)
                table_lines = []

        if not stripped_line:
            continue

        # 見出し
        if stripped_line.startswith('# '):
            writer.add_heading(stripped_line[2:], 1)
        elif stripped_line.startswith('## '):
            writer.add_heading(stripped_line[3:], 2)
        elif stripped_line.startswith('### '):
            writer.add_heading(stripped_line[4:], 3)
        elif stripped_line.startswith('#### '):
            writer.add_heading(stripped_line[5:], 4)
        
        # リスト
        elif stripped_line.startswith('* ') or stripped_line.startswith('- '):
            writer.add_paragraph(stripped_line[2:], style='List Bullet')
        elif stripped_line[0].isdigit() and stripped_line[1:3] == '. ':
            writer.add_paragraph(stripped_line[3:], style='List Number')
            
        # 引用
        elif stripped_line.startswith('> '):
            writer.add_quote(stripped_line[2:])
            
        # 画像
        elif stripped_line.startswith('![') and '](' in stripped_line and stripped_line.endswith(')'):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped_line)
            if match:
                image_path = match.group(2)
                if not os.path.isabs(image_path):
                    image_path = os.path.join(os.path.dirname(md_path), image_path)
                
                if os.path.exists(image_path):
                    writer.add_image(image_path, width_inches=5)
                else:
                    writer.add_paragraph(f"[画像が見つかりません: {image_path}]")

        # 通常の段落
        else:
            writer.add_paragraph(stripped_line)

    writer.save(output_path)

if __name__ == "__main__":
    import argparse
    import sys
    
    if len(sys.argv) < 2:
        print(f"""
Markdown to Docx / PDF Converter v{VERSION}
{CREDITS}

[Usage Examples]
  # Convert to Docx
  python3 md2docx.py input.md

  # Convert to PDF (Requires: sudo apt install python3-reportlab)
  python3 md2docx.py input.md --pdf
  python3 md2docx.py input.md output.pdf
""")
        sys.exit(1)
    
    parser = argparse.ArgumentParser(
        description=f'Convert Markdown to Docx or PDF. v{VERSION}\n{CREDITS}',
        formatter_class=argparse.RawTextHelpFormatter
    )
    parser.add_argument('--version', action='version', version=f'%(prog)s {VERSION}')
    parser.add_argument('input_file', help='Input Markdown file path')
    parser.add_argument('output_file', nargs='?', help='Output Docx or PDF file path')
    parser.add_argument('--pdf', action='store_true', help='Output as PDF (requires ReportLab)')
    
    args = parser.parse_args()

    input_file = args.input_file
    output_file = args.output_file
    to_pdf = args.pdf

    if not output_file:
        base, _ = os.path.splitext(input_file)
        if to_pdf:
            output_file = base + ".pdf"
        else:
            output_file = base + ".docx"
    else:
        # Check extensions implicitly
        if output_file.lower().endswith('.pdf'):
            to_pdf = True

    if os.path.exists(input_file):
        try:
            if to_pdf:
                convert_markdown(input_file, output_file, PdfWriter)
            else:
                convert_markdown(input_file, output_file, DocxWriter)
        except Exception as e:
            print(f"Error: {e}")
            sys.exit(1)
    else:
        print(f"File not found: {input_file}")
        sys.exit(1)
